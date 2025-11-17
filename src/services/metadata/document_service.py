from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class DocumentService:
    """ Service for extracting metadata from document files. """

    def __init__(self):
        self.office_formats = {'.docx', '.xlsx', '.pptx'}
        self.pdf_formats = {'.pdf'}
        self.text_formats = {'.txt', '.md', '.rtf'}
    
    def can_extract_metadata(self, file_path: Path) -> bool:
        """ Check if file is a supported document format. """
        suffix = file_path.suffix.lower()
        return (
            (suffix in self.office_formats and DOCX_AVAILABLE) or
            (suffix in self.pdf_formats and PDF_AVAILABLE) or
            suffix in self.text_formats
        )
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from document file
        Returns:
            Dict containing author, creation date, company, title, etc.
        """
        if not self.can_extract_metadata(file_path):
            return {}
        
        suffix = file_path.suffix.lower()
        try:
            if suffix == '.docx':
                return self._extract_docx_metadata(file_path)
            elif suffix == '.pdf':
                return self._extract_pdf_metadata(file_path)
            elif suffix in self.text_formats:
                return self._extract_text_metadata(file_path)
            else:
                return {}
        except Exception as e:
            print(f"Error extracting metadata from {file_path}: {e}")
            return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from Word document"""
        if not DOCX_AVAILABLE:
            return {}
        
        metadata = {}
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # Author information
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            
            # Company information
            if hasattr(core_props, 'company') and core_props.company:
                metadata['company'] = core_props.company
            
            # Title and subject
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.subject:
                metadata['subject'] = core_props.subject
            
            # Dates
            if core_props.created:
                metadata['created_date'] = core_props.created
            if core_props.modified:
                metadata['modified_date'] = core_props.modified
            
            # Document statistics
            if hasattr(core_props, 'category') and core_props.category:
                metadata['category'] = core_props.category
            
            # Keywords
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            
            # Comments
            if core_props.comments:
                metadata['comments'] = core_props.comments
            
            # Document content analysis
            word_count = len(doc.paragraphs)
            metadata['paragraph_count'] = word_count
            
            # Estimate reading time (assuming 200 words per minute)
            text_content = '\n'.join([para.text for para in doc.paragraphs])
            word_count = len(text_content.split())
            metadata['word_count'] = word_count
            metadata['estimated_reading_time'] = f"{max(1, word_count // 200)} minutes"
            
        except Exception as e:
            print(f"Error reading DOCX file {file_path}: {e}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF document"""
        if not PDF_AVAILABLE:
            return {}
        
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    
                    # Author information
                    if '/Author' in info:
                        metadata['author'] = str(info['/Author'])
                    if '/Creator' in info:
                        metadata['creator'] = str(info['/Creator'])
                    if '/Producer' in info:
                        metadata['producer'] = str(info['/Producer'])
                    
                    # Title and subject
                    if '/Title' in info:
                        metadata['title'] = str(info['/Title'])
                    if '/Subject' in info:
                        metadata['subject'] = str(info['/Subject'])
                    
                    # Dates
                    if '/CreationDate' in info:
                        metadata['creation_date'] = str(info['/CreationDate'])
                    if '/ModDate' in info:
                        metadata['modification_date'] = str(info['/ModDate'])
                    
                    # Keywords
                    if '/Keywords' in info:
                        metadata['keywords'] = str(info['/Keywords'])
                
                # Document statistics
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Estimate reading time (assuming 250 words per page)
                estimated_words = len(pdf_reader.pages) * 250
                metadata['estimated_word_count'] = estimated_words
                metadata['estimated_reading_time'] = f"{max(1, estimated_words // 200)} minutes"
        
        except Exception as e:
            print(f"Error reading PDF file {file_path}: {e}")
        
        return metadata
    
    def _extract_text_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic metadata from text files"""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                # Basic statistics
                lines = content.split('\n')
                words = content.split()
                
                metadata['line_count'] = len(lines)
                metadata['word_count'] = len(words)
                metadata['character_count'] = len(content)
                
                # Estimate reading time
                metadata['estimated_reading_time'] = f"{max(1, len(words) // 200)} minutes"
                
                # Try to detect if it's a specific format
                if file_path.suffix.lower() == '.md':
                    metadata['format'] = 'Markdown'
                    headers = [line for line in lines if line.startswith('#')]
                    metadata['header_count'] = len(headers)
                
        except Exception as e:
            print(f"Error reading text file {file_path}: {e}")
        
        return metadata
    
    def get_organization_date(self, file_path: Path) -> Optional[datetime]:
        """Get the best date for organizing this document"""
        metadata = self.extract_metadata(file_path)
        
        if 'created_date' in metadata and isinstance(metadata['created_date'], datetime):
            return metadata['created_date']
        
        if 'creation_date' in metadata:
            try:
                date_str = metadata['creation_date']
                if date_str.startswith('D:'):
                    date_str = date_str[2:16]  # Extract YYYYMMDDHHMMSS
                    return datetime.strptime(date_str, '%Y%m%d%H%M%S')
            except Exception:
                pass
        
        # Fallback to file modification time
        try:
            stat_result = file_path.stat()
            return datetime.fromtimestamp(stat_result.st_mtime)
        except OSError:
            return None
    
    def get_organization_info(self, file_path: Path) -> Dict[str, str]:
        """Get key information for file organization"""
        metadata = self.extract_metadata(file_path)
        
        org_info = {}
        
        # Author-based organization
        if 'author' in metadata:
            org_info['author'] = metadata['author']
        
        # Company-based organization
        if 'company' in metadata:
            org_info['company'] = metadata['company']
        
        # Category-based organization
        if 'category' in metadata:
            org_info['category'] = metadata['category']
        
        # Size-based organization for documents
        if 'word_count' in metadata:
            word_count = metadata['word_count']
            if word_count < 500:
                org_info['size_category'] = 'Short Documents'
            elif word_count < 2000:
                org_info['size_category'] = 'Medium Documents'
            else:
                org_info['size_category'] = 'Long Documents'
        
        return org_info
